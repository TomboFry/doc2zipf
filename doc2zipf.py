""" Outputs a list of words and their occurrences from a Word document """
from sys import argv
import operator
import re
from docx import Document

def gettext(filename):
    """ Return all the text contained inside a Word Document """
    doc = Document(filename)
    fulltext = []
    for paragraph in doc.paragraphs:
        fulltext.append(paragraph.text)

    zipf_file(' '.join(fulltext), filename)

def zipf_file(text, filename):
    """ Create a dictionary of each word and the number of times it appears in a string """

    # Remove all non-alphanumeric characters and make the text lowercase
    words = re.sub('[^0-9a-zA-Z ]+', '', text).lower().split()
    dictionary = {}

    # Go through every word and either add it to the dictionary, or increment
    # the number of occurrences
    for word in words:
        if word in dictionary:
            dictionary[word] = dictionary[word] + 1
        else:
            dictionary[word] = 1

    # Turn the dictionary into a list sorted by the number of occurrences
    sortedlist = sorted(dictionary.items(), key=operator.itemgetter(1), reverse=True)

    # Convert the list into a readable Word document
    list_to_file(sortedlist, filename)

def list_to_file(sorted_list, filename):
    """ Turns the sorted list into a Word document containing a table of results """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Word'
    hdr_cells[1].text = 'Occurrence'

    for key, value in sorted_list:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    doc.save("sorted - " + filename)

for arg in argv[1:]:
    gettext(arg)
